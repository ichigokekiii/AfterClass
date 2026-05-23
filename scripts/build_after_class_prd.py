from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)


@dataclass(frozen=True)
class Block:
    kind: str
    level: int | None = None
    text: str | None = None
    items: tuple[str, ...] | None = None


TITLE = "After Class Product Requirements Document"
SUBTITLE = "Build-ready MVP specification for a safety-first verified campus dating app"
DATE = "May 23, 2026"

METADATA = (
    ("Product", "After Class"),
    ("Document type", "Product requirements document"),
    ("Status", "Draft v1 for implementation planning"),
    ("Version", "1.0"),
    ("Last updated", DATE),
    ("Primary audience", "Founders, product, design, engineering, and future coding agents"),
)

BLOCKS: tuple[Block, ...] = (
    Block("heading", level=1, text="Executive Summary"),
    Block(
        "paragraph",
        text=(
            "After Class is a mobile-first dating app for verified college students that is designed to move users "
            "from discovery to a safer real-world meetup. Only active students age 18 or older with approved .edu "
            "identities may join. The MVP combines swipe-and-match discovery with university compatibility insights, "
            "a guided meetup engine, and a stronger safety layer that stays active before, during, and after the date."
        ),
    ),
    Block(
        "bullets",
        items=(
            "Trust layer: users may enter through manual .edu signup or Google OAuth, but the account used for access must itself be an approved .edu identity mapped to an activated school.",
            "Core loop: verify school identity, complete trust onboarding, choose Find People or Review a Friend, browse verified students, match, receive a meetup proposal, confirm arrival, complete the meetup, then optionally continue off-platform.",
            "Differentiator: the app translates internal university compatibility scores into human-friendly school-to-school labels, adds friend-reviewed social proof, and treats meetup safety as a core product behavior.",
            "Safety posture: exact home addresses are never required, location use is bounded to the meetup window with explicit per-date consent, and emergency flows prompt the user before escalation.",
        ),
    ),
    Block("heading", level=1, text="Product Vision and Principles"),
    Block(
        "paragraph",
        text=(
            "After Class should feel safer, more intentional, and less awkward than a generic dating app. "
            "The MVP exists to help active students discover each other, form mutual matches, meet in public daytime "
            "settings, and feel supported if something goes wrong."
        ),
    ),
    Block(
        "numbered",
        items=(
            "Verified identity first. School verification is part of the product promise, not a back-office check.",
            "Trust before browsing. Users complete the minimum identity and authenticity steps before entering discovery.",
            "Safety is product behavior. The app should actively support arrival, date progress, and post-date check-ins.",
            "Use data as a guide, not a claim of truth. Compatibility is directional and school-level, not a promise about individual chemistry.",
            "Keep v1 narrow and buildable. The MVP should prove verified student demand, meetup follow-through, and safety trustworthiness without becoming a full social platform.",
        ),
    ),
    Block("heading", level=1, text="MVP Scope"),
    Block("heading", level=2, text="In scope"),
    Block(
        "bullets",
        items=(
            "Student mobile onboarding with approved .edu verification through manual signup or Google OAuth using the same approved .edu identity.",
            "Mandatory face-to-profile verification before discovery access.",
            "A post-login chooser for Find People or Review a Friend when no profile exists.",
            "Profile setup, approximate area selection, travel radius selection, backup email capture, and trusted contact capture.",
            "Swipe-style discovery, mutual matching, and a meetup-first flow with no standard in-app one-to-one chat.",
            "A hybrid venue engine that prefers safety-approved venues in strong-coverage cities and can use online place data to fill coverage gaps.",
            "User-facing school compatibility insights shown as qualitative school-to-school labels.",
            "Real-time meetup safety flows, including arrival confirmation, meetup-window location support, shake-triggered emergency flow, and post-date safety check-ins.",
            "A post-meetup off-platform handoff flow for optional exchange of social accounts or contact methods.",
        ),
    ),
    Block("heading", level=2, text="Out of scope for MVP"),
    Block(
        "bullets",
        items=(
            "Paid subscriptions, boosts, or monetization systems.",
            "Voice notes, video calling, disappearing media, stories, or social feed mechanics.",
            "Student-facing web app experiences.",
            "Nighttime meetup scheduling, bar-first recommendations, or nightlife positioning.",
            "Always-on background tracking outside the explicit meetup window.",
            "A persistent in-app friends list, follow graph, or mutual-friends feature.",
            "Public reliability badges, public no-show indicators, or public safety ratings.",
        ),
    ),
    Block("heading", level=1, text="Core User Journeys"),
    Block("heading", level=2, text="Journey 1: Onboarding, verification, and trust setup"),
    Block(
        "numbered",
        items=(
            "User confirms they are at least 18 and chooses manual .edu signup or Google OAuth.",
            "The submitted or authenticated Google account must itself be an approved .edu identity tied to an activated school record.",
            "The app verifies school identity and then requires face-to-profile verification before discovery access. This is an authenticity check, not government-ID verification.",
            "The user provides profile basics, approximate area, travel radius, backup email, and a trusted contact.",
            "If no profile exists, the app presents the onboarding choices Find People and Review a Friend.",
        ),
    ),
    Block("heading", level=2, text="Journey 2: Review a friend"),
    Block(
        "numbered",
        items=(
            "If the user chooses Review a Friend, the app asks them to identify someone they know who already has an account.",
            "The app collects lightweight structured answers about that friend.",
            "Submitted answers become private-to-platform social-proof inputs and approved profile-facing review snippets for the reviewed user.",
            "Reviewer input may quietly influence ranking for the reviewed person, but reviewers cannot directly connect people in MVP.",
        ),
    ),
    Block("heading", level=2, text="Journey 3: Discovery and matching"),
    Block(
        "numbered",
        items=(
            "The user browses verified students in a swipe-style card flow.",
            "Cards are ranked by a hybrid score that combines explicit preference fit, distance fit, school compatibility, activity freshness, and safety eligibility.",
            "Mutual likes create a match and trigger the meetup proposal flow immediately.",
        ),
    ),
    Block("heading", level=2, text="Journey 4: Meetup planning and confirmation"),
    Block(
        "numbered",
        items=(
            "When a match is created, the app computes feasible meetup suggestions using both users' approximate areas, both travel radii, venue operating hours, and the daytime scheduling window.",
            "The app proposes one primary venue plus up to two alternates and three daytime slots between 8:00 AM and 6:00 PM local time.",
            "Venue sourcing is hybrid: safety-approved venues are preferred in strong-coverage cities, while online place data may supply additional candidate venues where coverage is sparse.",
            "If no safe valid venue exists, the app prompts users to widen radius or let the match expire.",
            "The meetup proposal screen must show a visible expiration countdown, and the match closes if no plan is accepted in time.",
        ),
    ),
    Block("heading", level=2, text="Journey 5: Arrival, meetup safety, and emergency support"),
    Block(
        "numbered",
        items=(
            "Before meetup-window safety features activate, the app asks for explicit per-date consent for location use during that meetup.",
            "When the meetup time arrives, the user is prompted to confirm arrival while proximity to the venue acts as supporting evidence.",
            "During the active meetup window, the app may use bounded location signals to detect likely arrival, sudden out-of-proximity events, or likely date completion.",
            "If the app detects a safety concern, it prompts the user first to confirm whether they are okay or need help.",
            "A shake gesture during the active meetup window opens an in-app emergency flow.",
            "If the user does not respond or signals danger, the app escalates to the trusted contact with a safety alert and relevant live or last-known location context.",
        ),
    ),
    Block("heading", level=2, text="Journey 6: Post-date feedback and continuation"),
    Block(
        "numbered",
        items=(
            "After the scheduled meetup time passes, a user confirms completion, or the system detects likely date-end behavior, the app prompts each user for adaptive post-date feedback.",
            "If a user reports that the other person did not show up, the app records a private internal no-show signal for trust, ranking, and moderation use only.",
            "After a completed meetup, the app may offer an optional off-platform continuation step where users choose whether to share supported social accounts or contact methods.",
            "For up to 24 hours after the meetup, the app runs repeated safety check-ins until the user confirms they are okay or the window ends per product policy.",
        ),
    ),
    Block("heading", level=1, text="UX Flow and Wireframes"),
    Block(
        "bullets",
        items=(
            "No standard in-app chat in MVP. The meetup proposal and safety flow are the center of the product.",
            "Safety is visible. Arrival confirmation, consent states, check-ins, and emergency actions should feel intentional rather than hidden.",
            "Identity should feel strong but not bureaucratic. School verification and face-to-profile verification should communicate trust, not surveillance.",
        ),
    ),
    Block("heading", level=2, text="Primary happy path"),
    Block(
        "bullets",
        items=(
            "Welcome, age gate, and auth choice.",
            "School verification and face check before discovery access.",
            "No-profile chooser, profile setup, area and radius, then safety setup.",
            "Discovery, match, meetup proposal, and per-date safety consent.",
            "Arrival check, active meetup safety tools, post-date feedback, and optional continue-elsewhere handoff.",
        ),
    ),
    Block("heading", level=2, text="Safety interaction rules"),
    Block(
        "bullets",
        items=(
            "Match just created — meetup proposal first; no normal chat thread.",
            "Meetup accepted — date details are pinned and per-date safety consent is requested.",
            "Meetup window begins — arrival prompt appears and safety tools become active.",
            "Suspicious movement or shake — user-first safety prompt or in-app emergency flow appears.",
            "Post-date complete — feedback and optional off-platform handoff appear.",
        ),
    ),
    Block("heading", level=1, text="Functional Requirements"),
    Block("heading", level=2, text="Verification and school identity"),
    Block(
        "bullets",
        items=(
            "FR-VER-1: The app must allow access only through approved .edu domains tied to active school records.",
            "FR-VER-2: The app must support manual .edu signup and Google OAuth using the same approved .edu identity.",
            "FR-VER-3: The Google account used for OAuth must itself be an approved .edu address.",
            "FR-VER-4: Verification must happen before the user can browse.",
            "FR-VER-5: Each approved school must have a canonical record and one or more approved email domains managed by admins.",
            "FR-VER-6: The system must classify the user to a school automatically from the verified domain.",
            "FR-VER-7: Accounts that lose verification status or are flagged as non-student misuse must be restricted pending review.",
            "FR-VER-8: The product must enforce a self-attested 18-plus age gate before account activation.",
            "FR-VER-9: Any verified user without a completed profile must see a chooser with Find People and Review a Friend before entering the main flow.",
        ),
    ),
    Block("heading", level=2, text="Profiles, trust setup, and preferences"),
    Block(
        "bullets",
        items=(
            "FR-PRO-1: Profiles must support display name, age, school, bio, interests, and three to six photos.",
            "FR-PRO-2: Users must provide an approximate area and a maximum travel radius in kilometers.",
            "FR-PRO-3: Exact addresses, dorm room details, and precise home coordinates must never be required profile fields.",
            "FR-PRO-4: Basic discovery preferences may include age range and gender interest, and smart likelihood ranking should rely primarily on explicit stated preferences in MVP.",
            "FR-PRO-5: Before discovery access, users must complete a face-to-profile verification step that confirms the onboarding face scan matches the profile photos.",
            "FR-PRO-6: Users must provide a backup email for recovery and safety messaging support.",
            "FR-PRO-7: Users must provide a separate trusted contact for emergency escalation flows.",
            "FR-PRO-8: If a profile includes approved friend-review content, the UI must show the reviewer's full name, school, and small profile photo beside the review snippet.",
        ),
    ),
    Block("heading", level=2, text="Friend review and social proof"),
    Block(
        "bullets",
        items=(
            "FR-SOC-1: The app must support a Review a Friend onboarding branch for verified users who know someone already on the platform.",
            "FR-SOC-2: The reviewed target must already have an account before a friend review can be submitted.",
            "FR-SOC-3: Friend reviews must become profile-level social proof for the reviewed user.",
            "FR-SOC-4: Reviewer input may influence ranking privately, but reviewers must not be able to directly connect people in MVP.",
            "FR-SOC-5: The platform must be able to moderate, approve, hide, or remove friend-review content.",
        ),
    ),
    Block("heading", level=2, text="Discovery, ranking, and prior-match memory"),
    Block(
        "bullets",
        items=(
            "FR-DIS-1: Discovery is cross-school by default across the currently approved school network.",
            "FR-DIS-2: The ranking model must combine explicit preference fit, distance fit, school compatibility, activity freshness, and safety eligibility.",
            "FR-DIS-3: School compatibility can influence ranking, but it must not fully block otherwise eligible users from appearing.",
            "FR-DIS-4: Distance fit must use approximate area centroids and user radius preferences rather than exact home locations.",
            "FR-DIS-5: Users who are blocked, banned, under review, or otherwise ineligible must never surface in discovery.",
            "FR-DIS-6: The system may store prior match and meetup history as internal memory, but previously matched pairs must be excluded from default rediscovery in MVP.",
        ),
    ),
    Block("heading", level=2, text="Matching"),
    Block(
        "bullets",
        items=(
            "FR-MAT-1: Mutual likes create a match immediately.",
            "FR-MAT-2: Match creation must emit analytics events for school-pair aggregation and post-match funnel tracking.",
            "FR-MAT-3: The first post-match screen must emphasize the meetup suggestion flow rather than a chat thread.",
            "FR-MAT-4: Every match must display a countdown until expiration on the meetup proposal screen.",
            "FR-MAT-5: When a match expires without a mutually accepted meetup plan, the match must move to an expired state and disappear from active match lists.",
        ),
    ),
    Block("heading", level=2, text="Meetup recommendation engine"),
    Block(
        "bullets",
        items=(
            "FR-MEET-1: Every match must trigger a meetup suggestion attempt unless one of the users is no longer eligible or safety-restricted.",
            "FR-MEET-2: The engine must use each user's approximate area centroid and declared radius to define the feasible meetup zone.",
            "FR-MEET-3: The engine must calculate a midpoint and search safe eligible venues that fit both users' radius limits.",
            "FR-MEET-4: Venue sourcing must be hybrid in MVP: prefer safety-approved venues in strong-coverage cities and allow online-data-assisted candidate sourcing where coverage is thin.",
            "FR-MEET-5: Bars, clubs, hotels, private residences, and other high-risk or low-visibility venues are excluded from the eligible set.",
            "FR-MEET-6: The engine must rank eligible venues by midpoint closeness, venue quality, daytime suitability, operating hours, and safety eligibility.",
            "FR-MEET-7: The app must propose one primary venue and up to two alternates when multiple valid options exist.",
            "FR-MEET-8: Suggested time slots must stay inside the MVP daytime window of 8:00 AM to 6:00 PM local time.",
            "FR-MEET-9: If no shared feasible venue exists, the app must clearly communicate that no valid suggestion is available and offer users the option to adjust radius or let the match expire.",
        ),
    ),
    Block("heading", level=2, text="Arrival, feedback, continuation, and no-show handling"),
    Block(
        "bullets",
        items=(
            "FR-POST-1: When the meetup window begins, the app must prompt each user to confirm arrival while using proximity as a supporting signal.",
            "FR-POST-2: After a scheduled meetup time passes, a user marks the meetup complete, or the system detects likely date-end behavior, the app must prompt each user for private adaptive feedback.",
            "FR-POST-3: If a user reports a no-show, the platform must record a private internal no-show signal for trust, ranking, and moderation use only.",
            "FR-POST-4: After a completed meetup, the app must offer an optional off-platform continuation flow where users may choose whether to share supported social accounts or contact methods.",
            "FR-POST-5: The MVP must not include a normal open-ended in-app one-to-one chat thread before or after match resolution.",
        ),
    ),
    Block("heading", level=2, text="School compatibility insights and data pipeline"),
    Block(
        "bullets",
        items=(
            "FR-COMP-1: The product must maintain school-to-school compatibility aggregates derived from recent platform behavior.",
            "FR-COMP-2: Inputs should include mutual-like conversion to match, meetup acceptance, meetup completion, no-show reporting, and other relevant qualitative or quantitative trust and outcome signals.",
            "FR-COMP-3: The default evaluation window is the trailing 90 days so the signal stays fresh.",
            "FR-COMP-4: Compatibility must be directional and school-level. It must not be framed as a person-to-person explanation of chemistry.",
            "FR-COMP-5: User-facing compatibility should be shown as qualitative labels such as High, Strong, Emerging, and Early Signal rather than raw scores.",
            "FR-COMP-6: School compatibility may influence discovery ranking, but it should contribute no more than a minority share of the final ranking score.",
            "FR-COMP-7: The platform must maintain a minimal clean data pipeline that normalizes qualitative and quantitative inputs for ranking, compatibility, trust, and future analytics.",
        ),
    ),
    Block("heading", level=2, text="Trust, safety, and privacy"),
    Block(
        "bullets",
        items=(
            "FR-SAFE-1: The product must include blocking and reporting on profiles and match flows. There is no normal chat surface in MVP.",
            "FR-SAFE-2: Meetup defaults must favor public venues and daytime hours only.",
            "FR-SAFE-3: Exact addresses, dorm names, and live-location history must not be exposed to other users in MVP.",
            "FR-SAFE-4: Meetup-window location monitoring must require explicit per-date consent before activation.",
            "FR-SAFE-5: Location monitoring must be bounded to the active meetup window and may not become always-on background tracking by default.",
            "FR-SAFE-6: During the active meetup window, the app must support a shake-triggered in-app emergency flow.",
            "FR-SAFE-7: If the app detects sudden out-of-proximity movement or likely date-end behavior, it must prompt the user first before escalating to anyone else.",
            "FR-SAFE-8: If the user does not respond or signals danger, the app must escalate to the trusted contact with a safety alert and relevant live or last-known location context.",
            "FR-SAFE-9: For up to 24 hours after the meetup, the app must support repeated safety check-ins until the user confirms they are okay or the window ends per policy.",
            "FR-SAFE-10: The app must support basic account restriction states such as active, under review, suspended, and banned.",
        ),
    ),
    Block("heading", level=2, text="Admin console"),
    Block(
        "bullets",
        items=(
            "FR-ADM-1: Admins must be able to create, edit, activate, and deactivate school records.",
            "FR-ADM-2: Admins must be able to add and review approved email domains per school.",
            "FR-ADM-3: Admins must be able to review queued reports, inspect relevant profile or meetup context, and take moderation actions.",
            "FR-ADM-4: Admins must be able to create, edit, deactivate, and safety-tag venues, including those sourced from external place data.",
            "FR-ADM-5: Admins must be able to inspect school compatibility aggregates and suppress misleading or low-quality school-pair displays if needed.",
            "FR-ADM-6: Admins must be able to review no-show patterns, safety escalations, and trusted-contact-triggered incidents.",
            "FR-ADM-7: Admin tooling must be internal-only and can be web-based even though the student product is mobile-first.",
        ),
    ),
    Block("heading", level=1, text="Product-side Domain Objects and Events"),
    Block("heading", level=2, text="Core domain objects"),
    Block(
        "bullets",
        items=(
            "School: canonical record for one college or university, including active status and metadata used in discovery and reporting.",
            "SchoolDomain: approved email domain mapped to a School record and managed by admins.",
            "User: verified student account with profile fields, trust state, preferences, approximate area, radius, backup email, and trusted contact metadata.",
            "FriendReview: attributed structured social-proof submission from one verified user about another user who already has an account.",
            "ApproxArea: user-selected place label plus stored centroid coordinates used in ranking and meetup generation.",
            "Venue: safe eligible meetup location with coordinates, category, operating hours, source type, and moderation or approval state.",
            "Match: record created when two users like each other, including timestamps and school-pair context.",
            "MeetupProposal: generated venue and slot recommendation tied to a match, with primary option, alternates, expiry, and user actions.",
            "MeetupSafetySession: per-date consent, arrival, proximity, check-in, and escalation context bounded to the meetup window and post-date follow-up.",
            "ContinuationHandoff: optional post-meetup exchange step for supported social accounts or contact methods.",
            "Report: trust-and-safety case attached to a user, profile, match, meetup, or escalation event.",
            "SchoolCompatibilityAggregate: rolling school-pair metrics and display band used in ranking and user-facing insight surfaces.",
        ),
    ),
    Block("heading", level=2, text="Required event tracking"),
    Block(
        "bullets",
        items=(
            "Account created, auth path selected, verification completed, face verification completed, and onboarding completed.",
            "No-profile chooser shown, Find People selected, Review a Friend selected, friend-review started, friend-review submitted, and friend-review approved or rejected.",
            "Discovery card shown, like, pass, mutual match created.",
            "Meetup proposal generated, proposal failed, venue accepted, alternate selected, slot changed, widen-radius requested, and suggestion expired.",
            "Safety consent granted or denied, arrival confirmed, shake emergency flow opened, user safety prompt shown, trusted-contact escalation triggered, and post-date safety check-in completed.",
            "Post-date feedback submitted, no-show reported, and continuation handoff accepted or skipped.",
            "Report filed, block applied, moderation action taken.",
            "School compatibility aggregate recalculated and school-pair display state changed.",
        ),
    ),
    Block("heading", level=1, text="Ranking, Compatibility, and Meetup Defaults"),
    Block(
        "bullets",
        items=(
            "Discovery ranking default weighting: preference fit and safety eligibility are mandatory gates; among eligible users, distance fit and activity freshness carry the largest weight, while school compatibility remains a meaningful but minority factor.",
            "Smart likelihood behavior: use explicit user preferences as the primary signal in MVP rather than opaque behavior-driven personalization.",
            "Compatibility default score bands: 80 to 100 High, 65 to 79 Strong, 50 to 64 Emerging, below threshold Early Signal rather than a punitive low label.",
            "Compatibility display thresholds: show a user-facing band only after at least 40 two-way likes and 20 matches for the school pair in the trailing 90 days.",
            "Meetup radius default: 10 km. Metric units are the only supported unit system in MVP.",
            "Time-slot default behavior: suggest three slots inside 8:00 AM to 6:00 PM local time, biased toward late morning, early afternoon, and late afternoon.",
            "Fallback behavior: first search near the midpoint, then search the broader overlap zone, then fail gracefully if no safe eligible venue satisfies both users' radius limits.",
        ),
    ),
    Block("heading", level=1, text="Success Metrics and Launch Criteria"),
    Block("heading", level=2, text="North-star outcome"),
    Block(
        "paragraph",
        text=(
            "The north-star outcome for MVP is a verified student match that progresses to a credible meetup plan and safely "
            "completes the date flow. The product should optimize for trust, meetup follow-through, and perceived safety "
            "rather than raw swipe volume alone."
        ),
    ),
    Block("heading", level=2, text="Core metrics"),
    Block(
        "bullets",
        items=(
            "Verification completion rate for eligible signups by auth path.",
            "Face-to-profile verification completion rate before discovery access.",
            "Profile completion rate before discovery access.",
            "No-profile chooser selection rate and friend-review submission rate.",
            "Like-to-match conversion rate.",
            "Percentage of matches that receive a valid meetup proposal within five seconds.",
            "Meetup proposal acceptance rate, alternate-selection rate, and expiration rate.",
            "Arrival confirmation rate and completed-meetup rate.",
            "Post-date feedback completion rate and no-show report rate.",
            "Safety check-in completion rate and trusted-contact escalation rate.",
            "Coverage of approved schools and safe eligible venues in the pilot geography.",
        ),
    ),
    Block("heading", level=2, text="Launch gates"),
    Block(
        "bullets",
        items=(
            "At least three approved schools are active in the initial pilot cluster.",
            "Each launch geography has enough approved venues and online-data-assisted coverage to support common midpoint suggestions across likely student corridors.",
            "Verification, face-to-profile verification, matching, meetup proposal generation, and safety check-ins work reliably enough that the core loop can be demonstrated end to end.",
            "Moderation workflows, trusted-contact escalation, and emergency deactivation are operational before public launch.",
        ),
    ),
    Block("heading", level=1, text="Risks and Mitigations"),
    Block(
        "bullets",
        items=(
            "Cold-start data risk: new schools will not have enough compatibility data. Mitigation: show Early Signal until thresholds are met and rely more heavily on distance and preference fit at launch.",
            "Venue coverage risk: midpoint math may fail in sparse or traffic-heavy areas. Mitigation: use hybrid venue sourcing with admin oversight in strong-coverage cities and external place data where coverage is thin.",
            "Privacy perception risk: location-aware safety features can feel invasive if poorly framed. Mitigation: require explicit per-date consent, bound location use to the meetup window, and never expose tracking to the other user.",
            "Verification loophole risk: some non-students may retain school mailboxes. Mitigation: state active-student-only policy, require face-to-profile authenticity checks, allow review states, and give admins tools to remove misuse.",
            "Product-friction risk: removing normal chat may feel unfamiliar to some users. Mitigation: keep the meetup flow fast, make post-date continuation optional, and validate whether the constrained flow improves real meetup completion.",
            "Social-proof abuse risk: users may submit low-quality, fake, or harmful friend reviews. Mitigation: require reviewed user account existence, keep moderation controls, and support hiding or removing review content quickly.",
        ),
    ),
    Block("heading", level=1, text="Acceptance Criteria for the Next Implementation Agent"),
    Block(
        "bullets",
        items=(
            "The implementation plan must preserve the active-student-only eligibility model and approved-domain school mapping.",
            "The student product must be mobile-first, while admin tooling may be a separate internal web surface.",
            "The core end-to-end flow must be verify school identity, complete face and safety onboarding, choose Find People or Review a Friend if no profile exists, browse, match, receive a meetup suggestion, confirm arrival, complete feedback, and optionally continue off-platform.",
            "Meetup generation must use approximate areas and radius preferences rather than exact address sharing or always-on live tracking.",
            "The MVP must exclude night meetups and high-risk venues.",
            "School compatibility must remain directional, school-to-school, threshold-gated, and qualitative on the user-facing side.",
            "Safety controls, moderation states, trusted-contact escalation, and meetup-window safety logic are required parts of MVP rather than post-launch extras.",
            "Friend-review social proof must replace any in-app friends or mutual-friends feature in MVP.",
            "The MVP must not reintroduce a normal one-to-one in-app chat thread without an explicit product decision.",
        ),
    ),
)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None, color=None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_border_bottom(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7D4E2")
    pbdr.append(bottom)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = BLACK
    title.paragraph_format.space_after = Pt(4)

    if "Subtitle" not in doc.styles:
        subtitle = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.italic = False
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run("After Class | Product Requirements Document")
    set_run_font(header_run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    footer_run = footer_p.add_run("Build-ready MVP spec")
    set_run_font(footer_run, size=9, color=MUTED)


def add_metadata(doc: Document) -> None:
    for label, value in METADATA:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=BLACK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border_bottom(rule)


def add_docx(docx_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.paragraph_format.space_after = Pt(2)
    run = intro.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph(TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(SUBTITLE, style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_metadata(doc)

    for block in BLOCKS:
        if block.kind == "heading":
            assert block.level is not None and block.text is not None
            doc.add_paragraph(block.text, style=f"Heading {block.level}")
        elif block.kind == "paragraph":
            assert block.text is not None
            doc.add_paragraph(block.text, style="Normal")
        elif block.kind == "bullets":
            assert block.items is not None
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.167
                p.add_run(item)
        elif block.kind == "numbered":
            assert block.items is not None
            for item in block.items:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.167
                p.add_run(item)
        else:
            raise ValueError(f"Unsupported block kind: {block.kind}")

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def add_markdown(md_path: Path) -> None:
    # Keep the checked-in Markdown file as the source of truth for PRD text.
    source_md = Path(__file__).resolve().parents[1] / "docs" / "after-class-prd.md"
    content = source_md.read_text(encoding="utf-8")
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(content, encoding="utf-8")


def add_pdf(pdf_path: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AfterClassTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=23,
        leading=28,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "AfterClassSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#666666"),
        spaceAfter=14,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#666666"),
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        textColor=colors.black,
        spaceAfter=6,
    )
    heading1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=16,
        spaceAfter=8,
    )
    heading2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=6,
    )
    heading3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "BulletBody",
        parent=body_style,
        spaceAfter=0,
        leftIndent=0,
    )

    story = [
        Paragraph("PRODUCT REQUIREMENTS DOCUMENT", label_style),
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, subtitle_style),
    ]

    for label, value in METADATA:
        story.append(Paragraph(f"<b>{label}:</b> {value}", body_style))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#C7D4E2")))
    story.append(Spacer(1, 10))

    heading_styles = {1: heading1, 2: heading2, 3: heading3}

    for block in BLOCKS:
        if block.kind == "heading":
            assert block.level is not None and block.text is not None
            story.append(Paragraph(block.text, heading_styles[block.level]))
        elif block.kind == "paragraph":
            assert block.text is not None
            story.append(Paragraph(block.text, body_style))
        elif block.kind == "bullets":
            assert block.items is not None
            items = [ListItem(Paragraph(item, bullet_style)) for item in block.items]
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=18,
                    bulletFontName="Helvetica",
                    bulletFontSize=10,
                    spaceAfter=8,
                )
            )
            story.append(Spacer(1, 4))
        elif block.kind == "numbered":
            assert block.items is not None
            items = [ListItem(Paragraph(item, bullet_style)) for item in block.items]
            story.append(
                ListFlowable(
                    items,
                    bulletType="1",
                    leftIndent=18,
                    bulletFontName="Helvetica",
                    bulletFontSize=10,
                    spaceAfter=8,
                )
            )
            story.append(Spacer(1, 4))
        else:
            raise ValueError(f"Unsupported block kind: {block.kind}")

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    def on_page(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(_doc.leftMargin, LETTER[1] - 0.6 * inch, "After Class | Product Requirements Document")
        footer_text = f"Build-ready MVP spec | Page {_doc.page}"
        canvas.drawRightString(LETTER[0] - _doc.rightMargin, 0.55 * inch, footer_text)
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build After Class PRD outputs.")
    parser.add_argument("--md", required=True, help="Path to the Markdown output.")
    parser.add_argument("--docx", required=True, help="Path to the DOCX output.")
    parser.add_argument("--pdf", help="Path to the PDF output.")
    args = parser.parse_args()

    md_path = Path(args.md)
    docx_path = Path(args.docx)

    add_markdown(md_path)
    add_docx(docx_path)
    if args.pdf:
        add_pdf(Path(args.pdf))


if __name__ == "__main__":
    main()
